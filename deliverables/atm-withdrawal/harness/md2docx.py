#!/usr/bin/env python3
"""將 Markdown 業務需求文件轉為 .docx(Step 2 產出物)。
支援:# 標題、清單(- / 1.)、表格(| a | b |)、粗體 **x**、段落。
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_runs(paragraph, text):
    # 處理 **bold**
    for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        run = paragraph.add_run(part)
        if i % 2 == 1:
            run.bold = True


def main(md_path, docx_path, title):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Noto Sans CJK TC"
    style.font.size = Pt(11)

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        # 表格
        if line.lstrip().startswith("|") and "|" in line[1:]:
            table_rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                row = lines[i].strip().strip("|")
                cells = [c.strip() for c in row.split("|")]
                if not re.match(r"^[\s:\-]+$", "".join(cells)):  # 跳過分隔列
                    table_rows.append(cells)
                i += 1
            if table_rows:
                cols = max(len(r) for r in table_rows)
                tbl = doc.add_table(rows=0, cols=cols)
                tbl.style = "Light Grid Accent 1"
                for r, cells in enumerate(table_rows):
                    crow = tbl.add_row().cells
                    for c in range(cols):
                        txt = cells[c] if c < len(cells) else ""
                        p = crow[c].paragraphs[0]
                        add_runs(p, txt)
                        if r == 0:
                            for run in p.runs:
                                run.bold = True
            continue
        # 標題
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2), level=level)
            i += 1
            continue
        # 清單
        m = re.match(r"^\s*[-*]\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, m.group(1))
            i += 1
            continue
        m = re.match(r"^\s*\d+\.\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(1))
            i += 1
            continue
        # 一般段落
        p = doc.add_paragraph()
        add_runs(p, line)
        i += 1

    doc.save(docx_path)
    print(f"wrote {docx_path}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2], sys.argv[3] if len(sys.argv) > 3 else "業務需求文件")
